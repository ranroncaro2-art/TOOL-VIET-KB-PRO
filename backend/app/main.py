import io
import json
import uuid
import asyncio
from datetime import datetime
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, Depends, HTTPException, BackgroundTasks, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.database import Base, engine, get_db, SessionLocal
from app.config import settings
import app.models as models
import app.schemas as schemas
from app.key_manager import KeyManager
from app.services.gemini_service import GeminiService
from app.services.monitor_service import MonitorService
from app.graphs.writing_graph import WritingGraph, WritingState
from app.graphs.validation_graph import ValidationGraph
from app.graphs.memory_graph import MemoryGraph

# Initialize FastAPI app
app = FastAPI(title=settings.PROJECT_NAME)

# CORS Middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory background task tracker
from app.services.queue_manager import QueueManager
queue_manager = QueueManager()
generation_tasks: Dict[str, Dict[str, Any]] = queue_manager.tasks

@app.on_event("startup")
async def startup_event():
    queue_manager.execute_callback = execute_queue_task
    queue_manager.start()
    
    # Create SQLite tables
    Base.metadata.create_all(bind=engine)
    
    # Seed default presets and prompts if empty
    db = next(get_db())
    try:
        # Seed Presets
        presets = [
            {"name": "Family Drama", "description": "High emotional tension, family secrets, domestic realism.", "variables": "{\"style\":\"realistic\",\"pov\":\"third_person\",\"target_length\":\"2000\"}"},
            {"name": "Manga Realistic", "description": "Manga-style pacing, detailed visual imagery, crisp dialogue.", "variables": "{\"style\":\"manga_realistic\",\"pov\":\"third_person\",\"target_length\":\"1800\"}"},
            {"name": "Dark Revenge", "description": "Suspenseful, Machiavellian strategies, dark internal thoughts.", "variables": "{\"style\":\"dark_psychological\",\"pov\":\"third_person\",\"target_length\":\"2500\"}"},
            {"name": "Romantic Comedy", "description": "Light-hearted, witty banter, slow-burn romantic tension.", "variables": "{\"style\":\"light_comedy\",\"pov\":\"third_person\",\"target_length\":\"2000\"}"}
        ]
        for p in presets:
            exists = db.query(models.PromptPreset).filter(models.PromptPreset.name == p["name"]).first()
            if not exists:
                db.add(models.PromptPreset(name=p["name"], description=p["description"], variables=p["variables"]))
        
        # Seed default prompts
        default_prompts = [
            # System Prompts (Cố định, priority 1)
            {
                "name": "Manga Screenwriter System", 
                "language": "ja", 
                "category": "System", 
                "is_cacheable": True, 
                "priority": 1,
                "content": "Bạn là biên kịch manga YouTube Nhật Bản chuyên tạo ra các câu chuyện có khả năng giữ chân người xem cao.\n\nNhiệm vụ:\n- Chuyển dàn ý thành kịch bản hoàn chỉnh.\n- Giữ logic thời gian và không gian.\n- Giữ nhất quán nhân vật.\n- Giữ nhất quán các chi tiết đã xuất hiện ở các chương trước.\n- Ưu tiên diễn biến, xung đột, cảm xúc và hành động.\n- Không tự ý thay đổi cốt truyện trong dàn ý.\n- Không giải thích quá trình suy luận.\n- Chỉ xuất nội dung truyện."
            },
            # Format Prompts (priority 2)
            {
                "name": "Manga Youtube Format", 
                "language": "ja", 
                "category": "Format", 
                "is_cacheable": True, 
                "priority": 2,
                "content": "[FORMAT TRÌNH BÀY KỊCH BẢN MANGA]\n- Lời dẫn viết tự nhiên.\n- Không dùng markdown.\n- Không dùng ký hiệu đặc biệt.\n- Lời thoại:\n  Nhân vật：「Lời thoại」\n- Độc thoại:\n  Nhân vật：（Nội tâm）\n- Tên riêng nhân vật viết bằng Katakana.\n- Nghiêm cấm sử dụng các dấu ngoặc kép dạng \"\" hoặc '' cho lời thoại. Phải ép buộc dùng dấu ngoặc vuông Nhật 「」 cho đối thoại và （） cho nội tâm.\n- Không viết: Nhân vật：「……」\n- Thay bằng: Nhân vật：「…ん」"
            },
            # Style Prompts (priority 3)
            {
                "name": "Manga Youtube JP Style", 
                "language": "ja", 
                "category": "Style", 
                "is_cacheable": True, 
                "priority": 3,
                "content": "[PHONG CÁCH MANGA YOUTUBE]\n- Ngôi 1.\n- Nhân vật chính xưng 俺.\n- Văn nói tự nhiên.\n- Không dùng văn học hàn lâm.\n- Nhịp kể vừa phải.\n- Show don't tell.\n- Tập trung hành động và tâm lý.\n- Đan xen nội tâm liên tục.\n- Tạo cảm giác nhập vai.\n- Ngôn ngữ phù hợp độ tuổi nhân vật."
            },
            # Default writers for other settings
            {
                "name": "Story Writer", 
                "language": "vi", 
                "category": "Style", 
                "is_cacheable": True, 
                "priority": 3,
                "content": "Bạn là nhà văn chuyên nghiệp. Viết bằng Tiếng Việt. Ngôi kể thứ ba. Văn phong tự nhiên, chân thực. Hãy tả chi tiết hành động và tâm lý nhân vật (Show, don't tell)."
            },
            {
                "name": "Story Writer", 
                "language": "en", 
                "category": "Style", 
                "is_cacheable": True, 
                "priority": 3,
                "content": "You are a professional novelist. Write in modern English. Third person narration. Use rich descriptions, psychological depth, and follow the 'Show, don't tell' rule."
            },
        ]
        for dp in default_prompts:
            exists = db.query(models.Prompt).filter(
                models.Prompt.name == dp["name"],
                models.Prompt.language == dp["language"]
            ).first()
            if not exists:
                p = models.Prompt(
                    name=dp["name"], 
                    language=dp["language"], 
                    category=dp["category"],
                    is_cacheable=dp["is_cacheable"],
                    priority=dp["priority"],
                    content=dp["content"]
                )
                db.add(p)
                db.flush()
                # Create initial version
                db.add(models.PromptVersion(prompt_id=p.id, version=1, content=dp["content"]))
                
        # Seed default system settings
        default_settings = [
            {"key": "api_mode", "value": "free", "description": "System API operation mode (free or paid)"},
            {"key": "workflow_mode", "value": "hybrid", "description": "System execution workflow mode (hybrid or flash_only)"},
            {"key": "free_mode_delay_seconds", "value": "4", "description": "Delay in seconds between LLM calls in Free Mode to avoid 429 rate limits (15 RPM)"},
            {"key": "model_pro", "value": "gemini-2.5-pro", "description": "Gemini Model assigned for Planner, Writer, Correction, and Deep Audits"},
            {"key": "model_flash", "value": "gemini-2.5-flash", "description": "Gemini Model assigned for Checkers, Summaries, Monitors, and Extractors"},
            {"key": "budget_max_input_tokens", "value": "6000", "description": "Strict limit for total context packer input tokens"},
            {"key": "budget_style_rules", "value": "500", "description": "Token quota allocated for Style Bible content"},
            {"key": "budget_selected_characters", "value": "1000", "description": "Token quota allocated for Character cast cores"},
            {"key": "budget_selected_world_rules", "value": "800", "description": "Token quota allocated for World Bible content"},
            {"key": "budget_current_arc", "value": "500", "description": "Token quota allocated for Current Arc Summary"},
            {"key": "budget_long_term_facts", "value": "300", "description": "Token quota allocated for Long-term Facts"},
            {"key": "budget_open_threads", "value": "300", "description": "Token quota allocated for Open Plot Threads"},
            {"key": "budget_retrieved_events", "value": "1000", "description": "Token quota allocated for Retrieved historical memory events"},
            {"key": "budget_outline", "value": "1000", "description": "Token quota allocated for Current chapter outline"}
        ]
        for ds in default_settings:
            exists = db.query(models.SystemSetting).filter(models.SystemSetting.key == ds["key"]).first()
            if not exists:
                db.add(models.SystemSetting(key=ds["key"], value=ds["value"], description=ds["description"]))

        db.commit()
    except Exception as e:
        print(f"Error seeding DB defaults: {str(e)}")
        db.rollback()
    finally:
        db.close()


# --- API KEYS ---
@app.get("/api/keys", response_model=List[schemas.ApiKeyResponse])
def list_keys(db: Session = Depends(get_db)):
    return db.query(models.ApiKey).order_by(models.ApiKey.created_at.desc()).all()

@app.post("/api/keys", response_model=schemas.ApiKeyResponse)
def add_key(data: schemas.ApiKeyCreate, db: Session = Depends(get_db)):
    try:
        new_key = KeyManager.register_key(db, data.key_value)
        return new_key
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.delete("/api/keys/{key_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_key(key_id: int, db: Session = Depends(get_db)):
    key = db.query(models.ApiKey).filter(models.ApiKey.id == key_id).first()
    if not key:
        raise HTTPException(status_code=404, detail="API Key not found")
    db.delete(key)
    db.commit()
    return None

@app.post("/api/keys/test")
def test_key(data: schemas.ApiKeyCreate):
    """Test a key value with a quick model check."""
    try:
        res = GeminiService.generate(
            api_key=data.key_value,
            model=settings.MODEL_FLASH,
            prompt="Respond with only 'OK'.",
            temperature=0.1
        )
        if "ok" in res.lower():
            return {"status": "success", "message": "API key is active and functional."}
        return {"status": "failed", "message": f"Unexpected API response: {res}"}
    except Exception as e:
        return {"status": "failed", "message": str(e)}


# --- STORIES & BIBLES ---
@app.get("/api/stories", response_model=List[schemas.StoryResponse])
def get_stories(db: Session = Depends(get_db)):
    return db.query(models.Story).order_by(models.Story.created_at.desc()).all()

@app.post("/api/stories", response_model=schemas.StoryResponse)
def create_story(data: schemas.StoryCreate, db: Session = Depends(get_db)):
    try:
        story = models.Story(
            title=data.title,
            description=data.description,
            language=data.language,
            pov=data.pov,
            style=data.style,
            target_length=data.target_length
        )
        db.add(story)
        db.flush()

        # Create initial World and Style Bibles
        # Prepopulate with dummy layout template
        db.add(models.StyleBible(
            story_id=story.id,
            version=1,
            content=f"POV: {data.pov}\nNarrative Style: {data.style}\nDialogue Rules: Natural and immersive\nCliffhanger Rules: Leave minor plot hook"
        ))
        
        db.add(models.WorldBible(
            story_id=story.id,
            version=1,
            content="Power System: Realistic modern day settings\nTimeline: Contemporary era\nKey Rules: Respect local cultural context"
        ))

        # Create first default arc
        db.add(models.Arc(
            story_id=story.id,
            arc_no=1,
            name="Introduction Arc",
            goal="Establish the setting and introduce principal characters",
            status="active"
        ))

        db.commit()
        db.refresh(story)
        return story
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=str(e))

@app.get("/api/stories/{story_id}", response_model=schemas.StoryDetailsResponse)
def get_story_details(story_id: int, db: Session = Depends(get_db)):
    story = db.query(models.Story).filter(models.Story.id == story_id).first()
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
        
    style_bible = db.query(models.StyleBible).filter(models.StyleBible.story_id == story_id).order_by(models.StyleBible.version.desc()).first()
    world_bible = db.query(models.WorldBible).filter(models.WorldBible.story_id == story_id).order_by(models.WorldBible.version.desc()).first()
    characters = db.query(models.Character).filter(models.Character.story_id == story_id).all()
    arcs = db.query(models.Arc).filter(models.Arc.story_id == story_id).order_by(models.Arc.arc_no).all()
    threads = db.query(models.PlotThread).filter(models.PlotThread.story_id == story_id).all()
    chapters = db.query(models.Chapter).filter(models.Chapter.story_id == story_id).order_by(models.Chapter.chapter_no).all()
    facts = db.query(models.LongTermFact).filter(models.LongTermFact.story_id == story_id).all()

    return {
        "story": story,
        "style_bible": style_bible,
        "world_bible": world_bible,
        "characters": characters,
        "arcs": arcs,
        "plot_threads": threads,
        "chapters": chapters,
        "long_term_facts": facts
    }

@app.delete("/api/stories/{story_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_story(story_id: int, db: Session = Depends(get_db)):
    story = db.query(models.Story).filter(models.Story.id == story_id).first()
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
    db.delete(story)
    db.commit()
    return None

def parse_story_outline(content: str, default_target: int = 2000) -> list:
    import re
    # Split content by [CHƯƠNG:
    chapters_raw = content.split("[CHƯƠNG:")
    parsed_chapters = []
    
    for chap_raw in chapters_raw:
        if not chap_raw.strip():
            continue
        
        # Extract the chapter number
        match_chap_no = re.match(r"\s*(\d+)\s*\]", chap_raw)
        if not match_chap_no:
            continue
        chapter_no = int(match_chap_no.group(1))
        
        # Extract title: [TIÊU_ĐỀ: ...]
        match_title = re.search(r"\[TIÊU_ĐỀ:\s*(.*?)\s*\]", chap_raw)
        title = match_title.group(1).strip() if match_title else f"Chương {chapter_no}"
        
        # Extract target: [TARGET: ...]
        match_target = re.search(r"\[TARGET:\s*(\d+)\s*\]", chap_raw)
        target = int(match_target.group(1)) if match_target else default_target
        
        # Split scenes by [CẢNH:
        scenes_raw = chap_raw.split("[CẢNH:")
        scene_plans = []
        
        for scene_raw in scenes_raw[1:]: # skip the part before the first scene
            match_scene_no = re.match(r"\s*(\d+)\s*\]", scene_raw)
            if not match_scene_no:
                continue
            scene_no = int(match_scene_no.group(1))
            
            # Extract content between [BẮT_ĐẦU_NỘI_DUNG] and [KẾT_THÚC_NỘI_DUNG]
            match_content = re.search(r"\[BẮT_ĐẦU_NỘI_DUNG\](.*?)\[KẾT_THÚC_NỘI_DUNG\]", scene_raw, re.DOTALL)
            scene_content = match_content.group(1).strip() if match_content else ""
            
            # Clean lines
            bullet_points = []
            for line in scene_content.splitlines():
                line_clean = line.strip()
                if line_clean:
                    bullet_points.append(line_clean)
            
            scene_plans.append({
                "scene_no": scene_no,
                "title": f"Scene {scene_no}",
                "characters_involved": [],
                "objective": "\n".join(bullet_points),
                "narrative_focus": "Outline Expansion"
            })
            
        parsed_chapters.append({
            "chapter_no": chapter_no,
            "title": title,
            "target_length": target,
            "scene_plan": scene_plans
        })
        
    return parsed_chapters

@app.post("/api/stories/{story_id}/import-outline")
def import_story_outline(story_id: int, data: schemas.ImportOutlineInput, db: Session = Depends(get_db)):
    story = db.query(models.Story).filter(models.Story.id == story_id).first()
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
        
    arc_id = data.arc_id
    if not arc_id:
        active_arc = db.query(models.Arc).filter(models.Arc.story_id == story_id, models.Arc.status == "active").first()
        if not active_arc:
            active_arc = models.Arc(
                story_id=story_id,
                arc_no=1,
                name="Default Arc",
                goal="Imported outline target arc",
                status="active"
            )
            db.add(active_arc)
            db.flush()
        arc_id = active_arc.id
        
    try:
        parsed_chapters = parse_story_outline(data.outline_text, default_target=story.target_length)
        if not parsed_chapters:
            raise ValueError("No valid chapters found in the outline text.")
            
        for chap_data in parsed_chapters:
            existing_chap = db.query(models.Chapter).filter(
                models.Chapter.story_id == story_id,
                models.Chapter.chapter_no == chap_data["chapter_no"]
            ).first()
            
            scene_plan_str = json.dumps(chap_data["scene_plan"])
            
            if existing_chap:
                existing_chap.title = chap_data["title"]
                existing_chap.target_length = chap_data["target_length"]
                existing_chap.scene_plan = scene_plan_str
            else:
                new_chap = models.Chapter(
                    story_id=story_id,
                    arc_id=arc_id,
                    chapter_no=chap_data["chapter_no"],
                    title=chap_data["title"],
                    target_length=chap_data["target_length"],
                    scene_plan=scene_plan_str
                )
                db.add(new_chap)
                
        db.commit()
        return {"status": "success", "imported_chapters_count": len(parsed_chapters)}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Failed to parse outline: {str(e)}")

@app.post("/api/stories/{story_id}/bibles/style")
def update_style_bible(story_id: int, content: Dict[str, str], db: Session = Depends(get_db)):
    last = db.query(models.StyleBible).filter(models.StyleBible.story_id == story_id).order_by(models.StyleBible.version.desc()).first()
    new_version = (last.version + 1) if last else 1
    new_bible = models.StyleBible(story_id=story_id, version=new_version, content=content.get("content", ""))
    db.add(new_bible)
    db.commit()
    return {"status": "success", "version": new_version}

@app.post("/api/stories/{story_id}/bibles/world")
def update_world_bible(story_id: int, content: Dict[str, str], db: Session = Depends(get_db)):
    last = db.query(models.WorldBible).filter(models.WorldBible.story_id == story_id).order_by(models.WorldBible.version.desc()).first()
    new_version = (last.version + 1) if last else 1
    new_bible = models.WorldBible(story_id=story_id, version=new_version, content=content.get("content", ""))
    db.add(new_bible)
    db.commit()
    return {"status": "success", "version": new_version}


# --- CHARACTERS ---
@app.post("/api/stories/{story_id}/characters", response_model=schemas.CharacterResponse)
def add_character(story_id: int, data: schemas.CharacterCreate, db: Session = Depends(get_db)):
    # Localize character name first if story language is Japanese and name is not in Japanese
    story = db.query(models.Story).filter(models.Story.id == story_id).first()
    name = data.name
    
    if story and story.language == "ja":
        # Check if name is simple latin and suggest ja name
        try:
            temp_key = KeyManager.get_utility_key(db)
            localized = GeminiService.localize_character_names(temp_key, [name], "ja")
            name = localized.get(name, name)
        except Exception:
            pass

    char = models.Character(
        story_id=story_id,
        name=name,
        profile_json="{}",
        appearance=data.appearance,
        personality=data.personality,
        speaking_style=data.speaking_style,
        secret=data.secret
    )
    db.add(char)
    db.commit()
    db.refresh(char)
    return char


# --- PROMPTS MANAGER ---
@app.get("/api/prompts", response_model=List[schemas.PromptResponse])
def get_prompts(db: Session = Depends(get_db)):
    return db.query(models.Prompt).all()

@app.post("/api/prompts")
def create_or_update_prompt(data: schemas.PromptCreate, db: Session = Depends(get_db)):
    # Prevent users from modifying System Prompts as they are fixed ("cố định")
    if data.category == "System":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN, 
            detail="System Prompts are fixed and cannot be modified by the user."
        )

    # Check if exists
    p = db.query(models.Prompt).filter(models.Prompt.name == data.name, models.Prompt.language == data.language).first()
    if p:
        p.content = data.content
        p.category = data.category
        p.is_cacheable = data.is_cacheable
        p.priority = data.priority
        # Find last version
        last_v = db.query(models.PromptVersion).filter(models.PromptVersion.prompt_id == p.id).order_by(models.PromptVersion.version.desc()).first()
        new_v = (last_v.version + 1) if last_v else 1
        db.add(models.PromptVersion(prompt_id=p.id, version=new_v, content=data.content))
        db.commit()
        return {"id": p.id, "version": new_v, "status": "updated"}
    else:
        new_p = models.Prompt(
            name=data.name, 
            language=data.language, 
            category=data.category,
            is_cacheable=data.is_cacheable,
            priority=data.priority,
            content=data.content
        )
        db.add(new_p)
        db.flush()
        db.add(models.PromptVersion(prompt_id=new_p.id, version=1, content=data.content))
        db.commit()
        return {"id": new_p.id, "version": 1, "status": "created"}

@app.get("/api/prompts/{prompt_id}/versions")
def get_prompt_versions(prompt_id: int, db: Session = Depends(get_db)):
    return db.query(models.PromptVersion).filter(models.PromptVersion.prompt_id == prompt_id).order_by(models.PromptVersion.version.desc()).all()

@app.post("/api/prompts/{prompt_id}/restore")
def restore_prompt_version(prompt_id: int, version_no: int, db: Session = Depends(get_db)):
    p = db.query(models.Prompt).filter(models.Prompt.id == prompt_id).first()
    v = db.query(models.PromptVersion).filter(models.PromptVersion.prompt_id == prompt_id, models.PromptVersion.version == version_no).first()
    if not p or not v:
        raise HTTPException(status_code=404, detail="Prompt or version not found")
        
    p.content = v.content
    # Create a new version listing to stack it
    last_v = db.query(models.PromptVersion).filter(models.PromptVersion.prompt_id == p.id).order_by(models.PromptVersion.version.desc()).first()
    db.add(models.PromptVersion(prompt_id=p.id, version=last_v.version + 1, content=v.content))
    db.commit()
    return {"status": "success", "restored_version": version_no}

@app.get("/api/presets", response_model=List[schemas.PromptPresetResponse])
def get_presets(db: Session = Depends(get_db)):
    presets = db.query(models.PromptPreset).all()
    # parse variables json string
    result = []
    for pr in presets:
        import json
        vars_dict = {}
        try:
            vars_dict = json.loads(pr.variables)
        except Exception:
            pass
        result.append(schemas.PromptPresetResponse(
            id=pr.id,
            name=pr.name,
            description=pr.description,
            variables=vars_dict
        ))
    return result


# --- CHAPTER GENERATION WORKER ---
async def chapter_generation_task(task_id: str, story_id: int, arc_id: int, chapter_no: int, title: str, outline: str, selected_characters: List[str], selected_world_rules: List[str]):
    # Setup database session
    db = SessionLocal()
    task = generation_tasks[task_id]
    
    try:
        task["logs"].append("Spinning up LangGraph writing engine...")
        
        # 1. Writing Graph
        w_state = WritingState(
            story_id=story_id,
            chapter_no=chapter_no,
            arc_id=arc_id,
            outline=outline,
            selected_characters=selected_characters,
            selected_world_rules=selected_world_rules,
            logs=task["logs"]
        )
        
        # Run graph
        w_state = await asyncio.to_thread(WritingGraph.run, db, w_state)
        task["draft"] = w_state.chapter_draft
        task["scene_plan"] = w_state.scene_plan
        task["logs"] = w_state.logs
        
        # 2. Validation Graph (Fact checker + style editor)
        task["logs"].append("Initiating parallel Validation checkers...")
        v_state = await ValidationGraph.validate_and_correct(db, w_state)
        task["draft"] = v_state.chapter_draft
        task["logs"] = v_state.logs
        
        # Save temp draft in task state
        task["status"] = "completed"
        task["logs"].append("Chapter generated successfully! Awaiting user approval.")
        
    except Exception as e:
        task["status"] = "failed"
        task["logs"].append(f"CRITICAL WORKFLOW EXCEPTION: {str(e)}")
    finally:
        db.close()

@app.post("/api/stories/{story_id}/chapters/generate")
async def start_chapter_generation(story_id: int, data: schemas.GenerateChapterInput, db: Session = Depends(get_db)):
    task_id = await queue_manager.add_task(
        task_type="chapter_generation",
        story_id=story_id,
        params={
            "arc_id": data.arc_id,
            "chapter_no": data.chapter_no,
            "title": data.title,
            "outline": data.outline,
            "selected_characters": data.selected_characters,
            "selected_world_rules": data.selected_world_rules
        }
    )
    return {"task_id": task_id, "status": "pending"}

@app.get("/api/tasks/{task_id}")
def get_task_status(task_id: str):
    task = generation_tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    return task

@app.get("/api/queue")
async def get_queue_list():
    return await queue_manager.get_queue_status()

@app.post("/api/queue/{task_id}/cancel")
async def cancel_queue_task(task_id: str):
    success = await queue_manager.cancel_task(task_id)
    if not success:
        raise HTTPException(status_code=404, detail="Task not found")
    return {"status": "success", "message": "Task cancelled successfully"}

@app.post("/api/stories/{story_id}/queue/cancel")
async def cancel_all_story_tasks(story_id: int):
    await queue_manager.cancel_story_tasks(story_id)
    return {"status": "success", "message": f"Cancelled all tasks for story {story_id}"}

async def story_auto_write_task(task_id: str, story_id: int, start_chap: int, end_chap: int, selected_characters: List[str], selected_world_rules: List[str]):
    db = SessionLocal()
    task = generation_tasks[task_id]
    
    try:
        for chap_no in range(start_chap, end_chap + 1):
            task["current_chapter"] = chap_no
            task["logs"].append(f"\n--- STARTING AUTOMATIC WRITING FOR CHAPTER {chap_no} ---")
            
            # Fetch chapter outline details from database
            chapter = db.query(models.Chapter).filter(
                models.Chapter.story_id == story_id,
                models.Chapter.chapter_no == chap_no
            ).first()
            
            if not chapter:
                task["logs"].append(f"Chapter {chap_no} record not found in database. Skipping...")
                continue
                
            outline_text = ""
            try:
                if chapter.scene_plan:
                    scenes = json.loads(chapter.scene_plan)
                    outline_text = " | ".join([s.get("objective", "") for s in scenes])
            except Exception:
                pass
                
            if not outline_text:
                outline_text = chapter.title
                
            task["logs"].append(f"Chapter {chap_no} title: '{chapter.title}'")
            task["logs"].append(f"Assembling context and running LangGraph writing engine...")
            
            # 1. Writing Graph
            w_state = WritingState(
                story_id=story_id,
                chapter_no=chap_no,
                arc_id=chapter.arc_id,
                outline=outline_text,
                selected_characters=selected_characters,
                selected_world_rules=selected_world_rules,
                logs=task["logs"]
            )
            
            w_state = await asyncio.to_thread(WritingGraph.run, db, w_state)
            
            # 2. Validation Graph
            task["logs"].append(f"Chapter {chap_no} draft completed. Running parallel checkers & correction...")
            v_state = await ValidationGraph.validate_and_correct(db, w_state)
            
            final_draft = v_state.chapter_draft
            
            # 3. Auto-Approve & Save
            task["logs"].append(f"Saving Chapter {chap_no} draft directly to database (Auto-Approve)...")
            chapter.content = final_draft
            db.commit()
            
            # 4. Synchronous Memory Compilation (to feed the next chapter)
            task["logs"].append(f"Compiling memory and entity graph changes...")
            MemoryGraph.run_all(
                db=db,
                story_id=story_id,
                arc_id=chapter.arc_id,
                chapter_no=chap_no,
                chapter_content=final_draft,
                logs=task["logs"]
            )
            
            # Arc Health and Monitor
            try:
                key = KeyManager.get_utility_key(db)
                health = MonitorService.monitor_arc_health(db, story_id, chapter.arc_id, key)
                if health.get("risk_level") == "HIGH":
                    task["logs"].append(f"WARNING: Arc Monitor flagged HIGH RISK: {health.get('reason')}")
            except Exception as h_err:
                pass
                
            # If Multiple of 10 chapters, compress arc
            if chap_no % 10 == 0:
                task["logs"].append("Multiple of 10 chapters reached. Compressing arc summary...")
                try:
                    key = KeyManager.get_utility_key(db)
                    MonitorService.compress_arc(db, story_id, chapter.arc_id, key)
                except Exception:
                    pass
            
            db.commit()
            task["logs"].append(f"✓ CHAPTER {chap_no} GENERATED & MEMORY COMPILED SUCCESSFULLY!")
            
            # 5. Delay to avoid rate limiting and allow DB commit cooling
            delay_seconds = int(models.SystemSetting.get_val(db, "free_mode_delay_seconds", "4"))
            task["logs"].append(f"Cooling down for {delay_seconds} seconds before starting Chapter {chap_no + 1}...")
            await asyncio.sleep(delay_seconds)
            
        task["status"] = "completed"
        task["logs"].append("🎉 ALL CHAPTERS AUTO-WRITTEN AND MEMORIES COMPILED SUCCESSFULLY!")
        
    except Exception as e:
        task["status"] = "failed"
        task["logs"].append(f"CRITICAL AUTO-WRITE TASK EXCEPTION: {str(e)}")
    finally:
        db.close()

@app.post("/api/stories/{story_id}/auto-write")
async def start_auto_write_loop(story_id: int, data: schemas.StartAutoWriteInput, db: Session = Depends(get_db)):
    story = db.query(models.Story).filter(models.Story.id == story_id).first()
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
        
    task_id = await queue_manager.add_task(
        task_type="auto_write",
        story_id=story_id,
        params={
            "start_chapter_no": data.start_chapter_no,
            "end_chapter_no": data.end_chapter_no,
            "selected_characters": data.selected_characters,
            "selected_world_rules": data.selected_world_rules
        }
    )
    
    return {"task_id": task_id, "status": "pending"}

async def execute_queue_task(task: Dict[str, Any]):
    task_type = task["type"]
    params = task["params"]
    story_id = task["story_id"]
    task_id = task["id"]
    
    if task_type == "auto_write":
        await story_auto_write_task(
            task_id=task_id,
            story_id=story_id,
            start_chap=params["start_chapter_no"],
            end_chap=params["end_chapter_no"],
            selected_characters=params["selected_characters"],
            selected_world_rules=params["selected_world_rules"]
        )
    else:
        await chapter_generation_task(
            task_id=task_id,
            story_id=story_id,
            arc_id=params["arc_id"],
            chapter_no=params["chapter_no"],
            title=params["title"],
            outline=params["outline"],
            selected_characters=params["selected_characters"],
            selected_world_rules=params["selected_world_rules"]
        )

async def approve_memory_compilation_task(story_id: int, arc_id: int, chapter_no: int, draft: str, task_id: str):
    db = SessionLocal()
    task = generation_tasks.get(task_id)
    # If the task object is already deleted, we can still run memory compilation but without updating task logs
    logs = task["logs"] if task else []
    try:
        logs.append("Compiling memory nodes (Event Extractor, Facts, Character deltas) in background...")
        MemoryGraph.run_all(
            db=db,
            story_id=story_id,
            arc_id=arc_id,
            chapter_no=chapter_no,
            chapter_content=draft,
            logs=logs
        )
        
        # Check Arc Health
        try:
            key = KeyManager.get_utility_key(db)
            health = MonitorService.monitor_arc_health(db, story_id, arc_id, key)
            if health.get("risk_level") == "HIGH":
                logs.append(f"WARNING: Arc Health Monitor flagged HIGH RISK. Reason: {health.get('reason')}")
        except Exception as health_err:
            logs.append(f"Arc health check failed: {str(health_err)}")

        # Trigger Arc Compression if chapter_no is a multiple of 10
        if chapter_no % 10 == 0:
            logs.append("Multiple of 10 chapters reached. Running Arc Compression summary...")
            try:
                key = KeyManager.get_utility_key(db)
                summary = MonitorService.compress_arc(db, story_id, arc_id, key)
                logs.append(f"Arc Compressed. 800-token Summary created: {summary[:100]}...")
            except Exception as comp_err:
                logs.append(f"Arc compression error: {str(comp_err)}")
                
        db.commit()
    except Exception as e:
        print(f"Error in background memory compilation: {str(e)}")
        if logs:
            logs.append(f"Error in background memory compilation: {str(e)}")
    finally:
        db.close()
        # Clean up task from memory after background work is finished
        # Commented out to preserve task execution logs in QueueManager history
        # if task_id in generation_tasks:
        #     del generation_tasks[task_id]
        pass

@app.post("/api/tasks/{task_id}/approve")
def approve_generated_chapter(task_id: str, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    task = generation_tasks.get(task_id)
    if not task or task["status"] != "completed":
        raise HTTPException(status_code=400, detail="Invalid task or task not completed.")
        
    story_id = task["story_id"]
    arc_id = task["arc_id"]
    chapter_no = task["chapter_no"]
    title = task["title"]
    draft = task["draft"]
    
    try:
        # Create chapter
        chapter = models.Chapter(
            story_id=story_id,
            arc_id=arc_id,
            chapter_no=chapter_no,
            title=title,
            content=draft,
            scene_plan=json.dumps(task.get("scene_plan", []))
        )
        db.add(chapter)
        db.commit()
        
        # Change status so frontend knows it's being compiled (if they poll)
        task["status"] = "approving"
        
        # Schedule memory compilation in background
        background_tasks.add_task(
            approve_memory_compilation_task,
            story_id=story_id,
            arc_id=arc_id,
            chapter_no=chapter_no,
            draft=draft,
            task_id=task_id
        )
        
        return {"status": "success", "chapter_id": chapter.id}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=f"Failed to approve chapter: {str(e)}")


# --- TRANSLATE STORY ---
@app.post("/api/stories/{story_id}/translate")
def translate_chapter(story_id: int, chapter_id: int, target_lang: str, db: Session = Depends(get_db)):
    chapter = db.query(models.Chapter).filter(models.Chapter.id == chapter_id, models.Chapter.story_id == story_id).first()
    story = db.query(models.Story).filter(models.Story.id == story_id).first()
    if not chapter or not story:
        raise HTTPException(status_code=404, detail="Chapter not found")
        
    try:
        key = KeyManager.get_utility_key(db)
        translated_text = GeminiService.translate_story(
            api_key=key,
            content=chapter.content,
            source_lang=story.language,
            target_lang=target_lang
        )
        return {"translated_content": translated_text}
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Translation failed: {str(e)}")


# --- EXPORT STORY ---
@app.get("/api/stories/{story_id}/export")
def export_story(story_id: int, format: str = "txt", db: Session = Depends(get_db)):
    import urllib.parse
    story = db.query(models.Story).filter(models.Story.id == story_id).first()
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
        
    chapters = db.query(models.Chapter).filter(models.Chapter.story_id == story_id).order_by(models.Chapter.chapter_no).all()
    
    if format == "txt":
        output = io.StringIO()
        output.write(f"=== {story.title} ===\n\n")
        output.write(f"Description: {story.description or ''}\n\n")
        for chap in chapters:
            output.write(f"\n\nChapter {chap.chapter_no}: {chap.title}\n")
            output.write(f"=========================================\n\n")
            output.write(chap.content or "")
        
        output.seek(0)
        safe_filename = urllib.parse.quote(story.title.replace(' ', '_'))
        return StreamingResponse(
            io.BytesIO(output.read().encode("utf-8")),
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}.txt"}
        )
        
    elif format == "docx":
        from docx import Document
        doc = Document()
        doc.add_heading(story.title, level=0)
        if story.description:
            doc.add_paragraph(story.description)
            
        for chap in chapters:
            doc.add_page_break()
            doc.add_heading(f"Chapter {chap.chapter_no}: {chap.title}", level=1)
            doc.add_paragraph(chap.content or "")
            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        safe_filename = urllib.parse.quote(story.title.replace(' ', '_'))
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}.docx"}
        )
        
    elif format == "pdf":
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        pdf_stream = io.BytesIO()
        doc = SimpleDocTemplate(pdf_stream, pagesize=letter)
        styles = getSampleStyleSheet()
        
        story_flow = []
        
        # Title
        story_flow.append(Paragraph(story.title, styles['Title']))
        story_flow.append(Spacer(1, 12))
        if story.description:
            story_flow.append(Paragraph(story.description, styles['Normal']))
            story_flow.append(Spacer(1, 24))
            
        # Clean HTML-like text for PDF flow
        for chap in chapters:
            story_flow.append(Paragraph(f"Chapter {chap.chapter_no}: {chap.title}", styles['Heading1']))
            story_flow.append(Spacer(1, 12))
            
            # Split text by paragraphs to make it flow beautifully
            content = chap.content or ""
            paragraphs = content.split("\n\n")
            for p in paragraphs:
                if p.strip():
                    # Escaping simple html chars
                    p_clean = p.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story_flow.append(Paragraph(p_clean, styles['Normal']))
                    story_flow.append(Spacer(1, 6))
            story_flow.append(Spacer(1, 24))
            
        doc.build(story_flow)
        pdf_stream.seek(0)
        safe_filename = urllib.parse.quote(story.title.replace(' ', '_'))
        return StreamingResponse(
            pdf_stream,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}.pdf"}
        )
        
    raise HTTPException(status_code=400, detail="Unsupported format. Choose txt, docx, or pdf.")


# --- SYSTEM SETTINGS ---
@app.get("/api/settings")
def get_system_settings(db: Session = Depends(get_db)):
    settings_list = db.query(models.SystemSetting).all()
    return {s.key: s.value for s in settings_list}

@app.post("/api/settings")
def update_system_settings(data: schemas.SystemSettingsBulkUpdate, db: Session = Depends(get_db)):
    try:
        for k, v in data.settings.items():
            setting = db.query(models.SystemSetting).filter(models.SystemSetting.key == k).first()
            if setting:
                setting.value = str(v)
            else:
                db.add(models.SystemSetting(key=k, value=str(v)))
        db.commit()
        return {"status": "success", "message": "System configurations updated successfully."}
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=400, detail=str(e))

# --- USAGE LOGS ---
@app.get("/api/stories/{story_id}/usage", response_model=schemas.UsageSummaryResponse)
def get_story_usage(story_id: int, db: Session = Depends(get_db)):
    story = db.query(models.Story).filter(models.Story.id == story_id).first()
    if not story:
        raise HTTPException(status_code=404, detail="Story not found")
        
    logs = db.query(models.UsageLog).filter(models.UsageLog.story_id == story_id).order_by(models.UsageLog.timestamp.desc()).all()
    
    total_cost = sum(l.estimated_cost for l in logs)
    total_input = sum(l.input_tokens for l in logs)
    total_output = sum(l.output_tokens for l in logs)
    total_cached = sum(l.cached_input_tokens for l in logs)
    
    # Calculate caching savings
    caching_savings = 0.0
    for l in logs:
        model_lower = l.model_name.lower()
        if "pro" in model_lower:
            caching_savings += l.cached_input_tokens * (0.00000125 - 0.0000003125)
        else:
            caching_savings += l.cached_input_tokens * (0.000000075 - 0.00000001875)
            
    return {
        "total_cost": total_cost,
        "total_input_tokens": total_input,
        "total_output_tokens": total_output,
        "total_cached_tokens": total_cached,
        "caching_savings": caching_savings,
        "logs": logs
    }
